# backend/ingestion/document_processor.py - Enhanced Document Processing
import hashlib
# import pymupdf
import fitz

from docx import Document
from typing import Tuple, Optional
import re, subprocess

class DocumentProcessor:
    """Enhanced document processing with improved text extraction"""
    
    @staticmethod
    def get_file_hash(file_bytes: bytes) -> str:
        """Generate hash for file deduplication"""
        return hashlib.md5(file_bytes).hexdigest()
    
    @staticmethod
    def extract_pdf_text(file_bytes: bytes) -> Tuple[str, dict]:
        """Extract text from PDF with metadata"""
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            text_parts = []
            metadata = {
                "page_count": doc.page_count,
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", ""),
                "subject": doc.metadata.get("subject", "")
            }
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                page_text = page.get_text()
                
                # Clean and process text
                page_text = DocumentProcessor._clean_text(page_text)
                if page_text.strip():
                    text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
            
            doc.close()
            full_text = "\n\n".join(text_parts)
            
            return full_text, metadata
            
        except Exception as e:
            raise Exception(f"PDF processing error: {str(e)}")
    
    @staticmethod
    def extract_docx_text(file_bytes: bytes) -> Tuple[str, dict]:
        """Extract text from DOCX with metadata"""
        try:
            doc = Document(file_bytes)
            text_parts = []
            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "title": "",
                "author": "",
                "subject": ""
            }
            
            # Extract core properties if available
            if hasattr(doc, 'core_properties'):
                metadata.update({
                    "title": doc.core_properties.title or "",
                    "author": doc.core_properties.author or "",
                    "subject": doc.core_properties.subject or ""
                })
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # Preserve structure with headers
                    if DocumentProcessor._is_header(text):
                        text_parts.append(f"\n### {text}\n")
                    else:
                        text_parts.append(text)
            
            full_text = "\n".join(text_parts)
            full_text = DocumentProcessor._clean_text(full_text)
            
            return full_text, metadata
            
        except Exception as e:
            raise Exception(f"DOCX processing error: {str(e)}")
    
    @staticmethod
    def _clean_text(text: str) -> str:
        """Clean and normalize text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Fix common OCR issues
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)  # Add space between camelCase
        text = re.sub(r'\.([A-Z])', r'. \1', text)  # Add space after periods
        
        # Remove page headers/footers patterns
        text = re.sub(r'^(Page \d+|\d+)$', '', text, flags=re.MULTILINE)
        
        # Normalize quotes
        text = re.sub(r'["""]', '"', text)
        # text = re.sub(r"[‘’]", "'", text)
        text = re.sub(r"[''']", "'", text)
        
        return text.strip()
    
    @staticmethod
    def _is_header(text: str) -> bool:
        """Detect if text is likely a header"""
        return (
            len(text) < 100 and
            text.isupper() or
            text.endswith(':') or
            re.match(r'^[\d\.\s]+[A-Z]', text)
        )

def extract_text_with_metadata(file_bytes: bytes, filename: str) -> Tuple[str, dict]:
    """Extract text and metadata from file based on extension"""
    ext = filename.lower().split('.')[-1]
    
    if ext == 'pdf':
        return DocumentProcessor.extract_pdf_text(file_bytes)
    elif ext in ['docx', 'doc']:
        return DocumentProcessor.extract_docx_text(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
